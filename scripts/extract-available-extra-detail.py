"""
Extract ONLY the Extra Detail files that are currently available (synced from OneDrive).
Updates the existing extracted area JSON files with the extra detail content.
"""
import json
from pathlib import Path
from docx import Document

EXTRA_DETAIL_DIR = Path(r"C:\Users\danrr\OneDrive - amazon.com\Desktop\CFI\_-_CFI_Lesson_Pack_from_George\Download All (PC) - CFI Lesson Pack\1. CFI Lessons & PPTs\3. CFI Extra Detail Lessons\CFI Extra Detail Individual Files - Nov25")
EXTRACTED_DIR = Path(r"C:\Users\danrr\OneDrive - amazon.com\Desktop\CFI\cfi-app-offline\scripts\extracted-content")

AREA_MAP = {
    "1. Fundamentals of Instructing": "I",
    "2. Technical Subject Areas": "II",
    "3. Preflight Preparation": "III",
    "4. Preflight Lesson on a Maneuver": "IV",
    "5. Preflight Procedures": "V",
    "6. Airport Operations": "VI",
    "7. Takeoffs, Landings, & Go-Arounds": "VII",
    "8. Fundamentals of Flight": "VIII",
    "9. Maneuvers": "IX",
    "10. Slow Flight, Stalls, & Spins": "X",
    "11. Basic Instrument Maneuvers": "XI",
    "12. Emergency Operations": "XII",
    "14. Postflight Procedures": "XIV",
    "15. Appendix": "APP",
}


def extract_docx_content(filepath):
    """Extract all text content from a .docx file."""
    try:
        doc = Document(str(filepath))
        content = {
            "filename": filepath.name,
            "paragraphs": [],
            "tables": [],
            "headings": [],
            "full_text": "",
        }
        full_text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            para_data = {
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "bold": any(run.bold for run in para.runs if run.bold),
            }
            content["paragraphs"].append(para_data)
            full_text_parts.append(text)
            if para.style and "Heading" in para.style.name:
                content["headings"].append(text)
        content["full_text"] = "\n".join(full_text_parts)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                content["tables"].append(table_data)
        return content
    except Exception as e:
        return {"filename": filepath.name, "error": str(e), "full_text": ""}


def main():
    print("=" * 60)
    print("EXTRA DETAIL EXTRACTION (Available Files Only)")
    print("=" * 60)

    total_extracted = 0
    total_failed = 0
    areas_updated = []

    for area_folder, area_roman in AREA_MAP.items():
        area_dir = EXTRA_DETAIL_DIR / area_folder
        if not area_dir.exists():
            continue

        docx_files = sorted(area_dir.glob("*.docx"))
        if not docx_files:
            continue

        # Test first file
        try:
            test_doc = Document(str(docx_files[0]))
            _ = len(test_doc.paragraphs)
        except Exception:
            print(f"  SKIP Area {area_roman} ({area_folder}) - files not synced")
            total_failed += len(docx_files)
            continue

        print(f"\n  Extracting Area {area_roman}: {len(docx_files)} files")

        # Load existing area JSON
        area_file = EXTRACTED_DIR / f"area_{area_roman}.json"
        if area_file.exists():
            with open(area_file, "r", encoding="utf-8") as f:
                area_data = json.load(f)
        else:
            area_data = {"area_number": area_roman, "area_name": area_folder, "standard_lessons": [], "extra_detail_lessons": [], "powerpoint_content": []}

        # Extract extra detail files
        extra_detail = []
        for docx_file in docx_files:
            try:
                content = extract_docx_content(docx_file)
                if content.get("full_text"):
                    extra_detail.append(content)
                    total_extracted += 1
                    print(f"    ✓ {docx_file.name} ({len(content['full_text'])} chars)")
                else:
                    print(f"    ✗ {docx_file.name} - empty or error: {content.get('error', 'no text')}")
                    total_failed += 1
            except Exception as e:
                print(f"    ✗ {docx_file.name} - {e}")
                total_failed += 1

        # Update area data
        area_data["extra_detail_lessons"] = extra_detail

        # Save updated area JSON
        with open(area_file, "w", encoding="utf-8") as f:
            json.dump(area_data, f, ensure_ascii=False, indent=2)
        print(f"    → Saved {len(extra_detail)} extra detail lessons to {area_file.name}")
        areas_updated.append(area_roman)

    print(f"\n{'='*60}")
    print(f"EXTRACTION COMPLETE")
    print(f"  Successfully extracted: {total_extracted}")
    print(f"  Failed/skipped: {total_failed}")
    print(f"  Areas updated: {', '.join(areas_updated)}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
